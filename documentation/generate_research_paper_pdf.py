# -*- coding: utf-8 -*-
"""
EduPulse AI - Research Paper PDF & DOCX Generator
Compiles documentation/EduPulse_AI_Research_Paper.md into a high-quality academic PDF.
"""

import os
import re
import sys
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable, Preformatted
)
from reportlab.pdfgen import canvas

BASE_DIR = r"d:\FINAL YEAR\EDUPULSE_AI_NEW"
DOC_DIR = os.path.join(BASE_DIR, "documentation")
MD_PATH = os.path.join(DOC_DIR, "EduPulse_AI_Research_Paper.md")
PDF_PATH = os.path.join(DOC_DIR, "EduPulse_AI_Research_Paper.pdf")
DOCX_PATH = os.path.join(DOC_DIR, "EduPulse_AI_Research_Paper.docx")

class NumberedCanvas(canvas.Canvas):
    """Two-pass canvas to dynamically compute and print total page count."""
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_header_footer(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_header_footer(self, page_count):
        self.saveState()
        self.setFont("Times-Roman", 8.5)
        self.setFillColor(colors.HexColor("#475569"))
        
        # Header (on pages after the first)
        if self._pageNumber > 1:
            self.drawString(54, 755, "EduPulse AI: Autonomous Behavioral Telemetry & Procrastination Mitigation")
            self.setStrokeColor(colors.HexColor("#cbd5e1"))
            self.setLineWidth(0.5)
            self.line(54, 747, 558, 747)
            
        # Footer
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 36, page_text)
        self.drawString(54, 36, "Research Paper — Department of Computer Science & Engineering")
        self.setStrokeColor(colors.HexColor("#cbd5e1"))
        self.setLineWidth(0.5)
        self.line(54, 48, 558, 48)
        
        self.restoreState()


def clean_markdown_inline(text):
    """Convert Markdown inline formatting to ReportLab XML tags."""
    # Escape XML entities first
    text = text.replace("&", "&amp;")
    # Bold **text**
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Italic *text*
    text = re.sub(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', r'<i>\1</i>', text)
    # Inline code `code`
    text = re.sub(r'`(.*?)`', r'<font face="Courier" color="#0f172a"><b>\1</b></font>', text)
    # Links [text](url) -> text (url)
    text = re.sub(r'\[(.*?)\]\((.*?)\)', r'<u>\1</u> (\2)', text)
    return text


def build_pdf():
    print(f"Reading Markdown from: {MD_PATH}")
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=A4,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()

    # Custom Typography Styles
    title_style = ParagraphStyle(
        'DocTitle',
        fontName='Times-Bold',
        fontSize=18,
        leading=22,
        alignment=1, # Center
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=10
    )

    author_style = ParagraphStyle(
        'DocAuthor',
        fontName='Times-Bold',
        fontSize=11,
        leading=15,
        alignment=1, # Center
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=3
    )

    affiliation_style = ParagraphStyle(
        'DocAffiliation',
        fontName='Times-Italic',
        fontSize=9.5,
        leading=13,
        alignment=1, # Center
        textColor=colors.HexColor("#475569"),
        spaceAfter=12
    )

    abstract_heading_style = ParagraphStyle(
        'AbstractHeading',
        fontName='Times-BoldItalic',
        fontSize=10.5,
        leading=13,
        textColor=colors.HexColor("#0f172a"),
        spaceBefore=6,
        spaceAfter=4
    )

    abstract_body_style = ParagraphStyle(
        'AbstractBody',
        fontName='Times-Roman',
        fontSize=9.5,
        leading=13.5,
        alignment=4, # Justified
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=6
    )

    h1_style = ParagraphStyle(
        'Heading1_Custom',
        fontName='Times-Bold',
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#0f172a"),
        spaceBefore=16,
        spaceAfter=6,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'Heading2_Custom',
        fontName='Times-Bold',
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#1e293b"),
        spaceBefore=12,
        spaceAfter=4,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'Heading3_Custom',
        fontName='Times-BoldItalic',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#334155"),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'Body_Custom',
        fontName='Times-Roman',
        fontSize=9.5,
        leading=13.5,
        alignment=4, # Justified
        textColor=colors.HexColor("#111827"),
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'Bullet_Custom',
        fontName='Times-Roman',
        fontSize=9.5,
        leading=13.5,
        alignment=4,
        textColor=colors.HexColor("#111827"),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    code_style = ParagraphStyle(
        'Code_Custom',
        fontName='Courier',
        fontSize=7.5,
        leading=9.5,
        textColor=colors.HexColor("#0f172a")
    )

    table_header_style = ParagraphStyle(
        'TableHeader',
        fontName='Times-Bold',
        fontSize=8.5,
        leading=10.5,
        textColor=colors.HexColor("#ffffff"),
        alignment=1 # Center
    )

    table_cell_style = ParagraphStyle(
        'TableCell',
        fontName='Times-Roman',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#0f172a")
    )

    table_cell_bold = ParagraphStyle(
        'TableCellBold',
        fontName='Times-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#0f172a")
    )

    story = []

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        line_s = line.strip()

        # Handle Code Blocks
        if line_s.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = "\n".join(code_lines)
                table_code = Table([[Preformatted(code_text, code_style)]], colWidths=[500])
                table_code.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                    ('BOX', (0, 0), (-1, -1), 0.75, colors.HexColor("#cbd5e1")),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ]))
                story.append(Spacer(1, 4))
                story.append(table_code)
                story.append(Spacer(1, 6))
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle Markdown Tables
        if line_s.startswith('|') and line_s.endswith('|'):
            table_lines.append(line_s)
            i += 1
            # Check if next line is not table
            if i >= len(lines) or not (lines[i].strip().startswith('|') and lines[i].strip().endswith('|')):
                # Process table
                parsed_rows = []
                for t_idx, t_line in enumerate(table_lines):
                    # Skip separator line |---|---|
                    if re.match(r'^\|(\s*:?-+:?\s*\|)+$', t_line):
                        continue
                    raw_cells = [c.strip() for c in t_line.split('|')[1:-1]]
                    row_cells = []
                    for c_idx, cell in enumerate(raw_cells):
                        cell_clean = clean_markdown_inline(cell)
                        if t_idx == 0:
                            row_cells.append(Paragraph(cell_clean, table_header_style))
                        else:
                            if cell.startswith('**') and cell.endswith('**'):
                                row_cells.append(Paragraph(cell_clean, table_cell_bold))
                            else:
                                row_cells.append(Paragraph(cell_clean, table_cell_style))
                    if row_cells:
                        parsed_rows.append(row_cells)

                if parsed_rows:
                    num_cols = max(len(r) for r in parsed_rows)
                    # Adjust column widths
                    if num_cols == 2:
                        widths = [150, 350]
                    elif num_cols == 3:
                        widths = [140, 160, 200]
                    elif num_cols == 4:
                        widths = [120, 110, 110, 160]
                    elif num_cols == 5:
                        widths = [110, 95, 95, 100, 100]
                    elif num_cols == 6:
                        widths = [110, 80, 75, 75, 80, 80]
                    elif num_cols == 7:
                        widths = [95, 65, 65, 65, 65, 75, 70]
                    elif num_cols == 8:
                        widths = [95, 55, 55, 55, 55, 60, 60, 65]
                    elif num_cols == 9:
                        widths = [85, 50, 50, 50, 50, 50, 55, 55, 55]
                    else:
                        col_w = 500 / num_cols
                        widths = [col_w] * num_cols

                    t = Table(parsed_rows, colWidths=widths[:num_cols], repeatRows=1)
                    t_style = [
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                        ('LEFTPADDING', (0, 0), (-1, -1), 5),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ]
                    # Alternating row colors
                    for r_i in range(1, len(parsed_rows)):
                        if r_i % 2 == 0:
                            t_style.append(('BACKGROUND', (0, r_i), (-1, r_i), colors.HexColor("#f8fafc")))
                    t.setStyle(TableStyle(t_style))
                    story.append(Spacer(1, 4))
                    story.append(t)
                    story.append(Spacer(1, 6))

                table_lines = []
            continue

        # Skip empty lines
        if not line_s:
            i += 1
            continue

        # Document Title
        if line_s.startswith('# '):
            title_text = clean_markdown_inline(line_s[2:].strip())
            story.append(Paragraph(title_text, title_style))
            i += 1
            continue

        # Author / Affiliation lines at the top
        if line_s.startswith('**Abdul Samad**') or line_s.startswith('*Department of'):
            p_text = clean_markdown_inline(line_s)
            if 'Abdul Samad' in line_s:
                story.append(Paragraph(p_text, author_style))
            else:
                story.append(Paragraph(p_text, affiliation_style))
            i += 1
            continue

        if line_s.startswith('*GitHub:') or line_s.startswith('*Email:'):
            p_text = clean_markdown_inline(line_s)
            story.append(Paragraph(p_text, affiliation_style))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cbd5e1"), spaceAfter=10))
            i += 1
            continue

        # Horizontal Rule
        if line_s == '---':
            story.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor("#cbd5e1"), spaceBefore=6, spaceAfter=8))
            i += 1
            continue

        # Abstract
        if line_s.startswith('### Abstract'):
            story.append(Paragraph("<b>ABSTRACT</b>", abstract_heading_style))
            i += 1
            continue

        # Section Headings
        if line_s.startswith('## '):
            h1_text = clean_markdown_inline(line_s[3:].strip())
            story.append(Paragraph(h1_text, h1_style))
            i += 1
            continue

        if line_s.startswith('### '):
            h2_text = clean_markdown_inline(line_s[4:].strip())
            story.append(Paragraph(h2_text, h2_style))
            i += 1
            continue

        if line_s.startswith('#### '):
            h3_text = clean_markdown_inline(line_s[5:].strip())
            story.append(Paragraph(h3_text, h3_style))
            i += 1
            continue

        # Bullet points
        if line_s.startswith('- ') or line_s.startswith('* '):
            b_text = clean_markdown_inline(line_s[2:].strip())
            bullet_p = Paragraph(f"&bull; {b_text}", bullet_style)
            story.append(bullet_p)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)$', line_s)
        if num_match:
            n_num = num_match.group(1)
            n_text = clean_markdown_inline(num_match.group(2))
            story.append(Paragraph(f"<b>{n_num}.</b> {n_text}", bullet_style))
            i += 1
            continue

        # Math equations display ($$ ... $$)
        if line_s.startswith('$$') and line_s.endswith('$$') and len(line_s) > 4:
            math_content = line_s[2:-2].strip()
            # Format display math in a clean shaded box
            math_p = Paragraph(f"<i>{math_content}</i>", ParagraphStyle(
                'MathBox', fontName='Times-Italic', fontSize=10, leading=14, alignment=1, textColor=colors.HexColor("#0f172a")
            ))
            t_math = Table([[math_p]], colWidths=[500])
            t_math.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ]))
            story.append(Spacer(1, 3))
            story.append(t_math)
            story.append(Spacer(1, 5))
            i += 1
            continue

        # Standard Paragraph
        p_clean = clean_markdown_inline(line_s)
        # Keywords
        if p_clean.startswith('<b>Keywords:</b>') or p_clean.startswith('Keywords:'):
            story.append(Paragraph(f"<i>{p_clean}</i>", abstract_body_style))
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(p_clean, body_style))

        i += 1

    print("Building PDF with ReportLab NumberedCanvas...")
    doc.build(story, canvasmaker=NumberedCanvas)
    pdf_size = os.path.getsize(PDF_PATH)
    print(f"[SUCCESS] PDF Generated: {PDF_PATH} ({pdf_size:,} bytes)")


def build_docx():
    """Also generates a clean formatted Word DOCX file for academic submission."""
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        print(f"Building DOCX version at: {DOCX_PATH}")
        doc = Document()

        for s in doc.sections:
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)

        normal = doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

        with open(MD_PATH, "r", encoding="utf-8") as f:
            lines = f.readlines()

        for line in lines:
            line_s = line.strip()
            if not line_s:
                continue
            if line_s.startswith('# '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line_s[2:].strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif line_s.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(line_s[3:].strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif line_s.startswith('### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(line_s[4:].strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            elif line_s.startswith('- ') or line_s.startswith('* '):
                p = doc.add_paragraph(line_s[2:].strip(), style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
            else:
                p = doc.add_paragraph(line_s)
                p.paragraph_format.space_after = Pt(4)

        doc.save(DOCX_PATH)
        print(f"[SUCCESS] DOCX Generated: {DOCX_PATH} ({os.path.getsize(DOCX_PATH):,} bytes)")
    except Exception as e:
        print(f"Notice: DOCX compilation skipped: {e}")

if __name__ == '__main__':
    build_pdf()
    build_docx()
