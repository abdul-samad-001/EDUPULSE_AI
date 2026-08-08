"""
train_model.py

EduPulse AI - Productivity Prediction ML Training & Evaluation Pipeline

Train multiple regression models (Linear Regression, Random Forest Regressor,
XGBoost / Gradient Boosting Regressor) using the synthetic productivity dataset.
Generates comprehensive evaluation metrics (MAE, MSE, RMSE, R2, Explained Variance),
prediction scatter plots, residual plots, feature importances, learning curves,
saves the best model & metadata, and generates PDF & DOCX report documents.
"""

import os
import sys
import json
import time
import joblib
import numpy as np
import pandas as pd
from pathlib import Path
from datetime import datetime

# Set stdout/stderr encoding to utf-8 if possible
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

import matplotlib
matplotlib.use("Agg")  # Non-interactive backend for headless plot generation
import matplotlib.pyplot as plt

from sklearn.model_selection import train_test_split, learning_curve
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor

# Try importing XGBoost, fallback to GradientBoostingRegressor if unavailable
try:
    # pyrefly: ignore [missing-import]
    from xgboost import XGBRegressor
    HAS_XGB = True
except ImportError:
    from sklearn.ensemble import GradientBoostingRegressor
    HAS_XGB = False

from sklearn.metrics import (
    mean_absolute_error,
    mean_squared_error,
    r2_score,
    explained_variance_score,
)

# ==========================================================
# STEP 12 & REQUIREMENT: DIRECTORY CHECK & PATH SETUP
# ==========================================================

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

if SCRIPT_DIR.name == "productivity":
    BASE_DIR = SCRIPT_DIR.parent.parent
elif SCRIPT_DIR.name == "scripts":
    BASE_DIR = SCRIPT_DIR.parent
else:
    BASE_DIR = Path.cwd()

DATA_DIR = BASE_DIR / "data" / "productivity"
MODEL_DIR = BASE_DIR / "models" / "productivity"
EVAL_DIR = BASE_DIR / "evaluation" / "productivity"
REPORTS_DIR = BASE_DIR / "reports"

DATA_DIR.mkdir(parents=True, exist_ok=True)
MODEL_DIR.mkdir(parents=True, exist_ok=True)
EVAL_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# Start global timing
pipeline_start_time = time.time()


# ==========================================================
# STEP 1 - Load Dataset
# ==========================================================

print("=" * 65)
print("STEP 1: Loading Productivity Dataset...")
print("=" * 65)

dataset_path = DATA_DIR / "productivity_dataset.csv"
if not dataset_path.exists():
    dataset_path = BASE_DIR / "data" / "productivity" / "productivity_dataset.csv"

if not dataset_path.exists():
    raise FileNotFoundError(
        f"Dataset not found at '{dataset_path}'. Please run generate_dataset.py first."
    )

df = pd.read_csv(dataset_path)

print("Dataset Loaded Successfully")
print(f"Dataset Path: {dataset_path}")
print(f"Dataset Shape: {df.shape}")

print("\nFirst 5 Records:")
print(df.head())

feature_names = [col for col in df.columns if col != "productivity_score"]
print(f"\nFeature Names ({len(feature_names)}):")
for feat in feature_names:
    print(f" [+] {feat}")


# ==========================================================
# STEP 2 - Separate Features and Target
# ==========================================================

print("\n" + "=" * 65)
print("STEP 2: Preparing Features and Target...")
print("=" * 65)

X = df[feature_names]
y = df["productivity_score"]

print(f"Features Shape : {X.shape}")
print(f"Target Shape   : {y.shape}")
print("Target Variable: productivity_score (Range: 0.00 to 100.00)")


# ==========================================================
# STEP 3 - Train/Test Split
# ==========================================================

print("\n" + "=" * 65)
print("STEP 3: Splitting Dataset & Saving Train/Test Splits...")
print("=" * 65)

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.20, random_state=42
)

print(f"Training Samples : {len(X_train):,}")
print(f"Testing Samples  : {len(X_test):,}")
print(f"X_train Shape    : {X_train.shape}")
print(f"X_test Shape     : {X_test.shape}")

# Save train and test sets to CSV
train_df = pd.concat([X_train, y_train], axis=1)
test_df = pd.concat([X_test, y_test], axis=1)

train_path = DATA_DIR / "train.csv"
test_path = DATA_DIR / "test.csv"

train_df.to_csv(train_path, index=False)
test_df.to_csv(test_path, index=False)

print(f"[+] train.csv saved -> {train_path}")
print(f"[+] test.csv saved  -> {test_path}")


# ==========================================================
# STEP 4 - Feature Scaling
# ==========================================================

print("\n" + "=" * 65)
print("STEP 4: Scaling Features...")
print("=" * 65)

scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train)
X_test_scaled = scaler.transform(X_test)

scaler_path = MODEL_DIR / "scaler.pkl"
joblib.dump(scaler, scaler_path)

print("Feature Scaling Completed")
print(f"Scaled Training Shape : {X_train_scaled.shape}")
print(f"Scaled Testing Shape  : {X_test_scaled.shape}")
print(f"[+] Scaler saved      -> {scaler_path}")


# ==========================================================
# STEP 5 & 6 - Model Training & Evaluation
# ==========================================================

print("\n" + "=" * 65)
print("STEP 5 & 6: Training and Evaluating Models...")
print("=" * 65)

# Define models list
models_to_train = [
    ("Linear Regression", LinearRegression(), False),
    ("Random Forest Regressor", RandomForestRegressor(n_estimators=100, random_state=42, n_jobs=-1), True),
]

if HAS_XGB:
    models_to_train.append(
        ("XGBoost Regressor", XGBRegressor(n_estimators=100, learning_rate=0.1, random_state=42, n_jobs=-1), True)
    )
else:
    models_to_train.append(
        ("Gradient Boosting Regressor", GradientBoostingRegressor(n_estimators=100, learning_rate=0.1, random_state=42), True)
    )

results_list = []
trained_models_dict = {}

print("Training Started...\n")

for idx, (name, model, uses_raw_features) in enumerate(models_to_train, 1):
    print(f"[{idx}/{len(models_to_train)}] Training {name}...")
    t0 = time.time()
    
    # Train on raw or scaled features depending on algorithm
    X_tr = X_train_scaled if not uses_raw_features else X_train
    X_te = X_test_scaled if not uses_raw_features else X_test
    
    model.fit(X_tr, y_train)
    t1 = time.time()
    train_duration = t1 - t0
    
    # Predict
    y_pred = model.predict(X_te)
    
    # Metrics
    mae = mean_absolute_error(y_test, y_pred)
    mse = mean_squared_error(y_test, y_pred)
    rmse = np.sqrt(mse)
    r2 = r2_score(y_test, y_pred)
    exp_var = explained_variance_score(y_test, y_pred)
    
    res = {
        "model_name": name,
        "model_obj": model,
        "uses_raw_features": uses_raw_features,
        "mae": float(mae),
        "mse": float(mse),
        "rmse": float(rmse),
        "r2": float(r2),
        "explained_variance": float(exp_var),
        "training_time_sec": float(train_duration),
        "y_pred": y_pred,
    }
    
    results_list.append(res)
    trained_models_dict[name] = model
    
    print(f"{name} Training Completed ({train_duration:.2f}s)")
    print(f"  MAE                 : {mae:.4f}")
    print(f"  MSE                 : {mse:.4f}")
    print(f"  RMSE                : {rmse:.4f}")
    print(f"  R² Score            : {r2:.4f}")
    print(f"  Explained Variance  : {exp_var:.4f}\n")

print("Training Completed.")


# ==========================================================
# STEP 7 - Model Comparison & Selection
# ==========================================================

print("=" * 65)
print("STEP 7: Model Comparison & Selection...")
print("=" * 65)

comp_df = pd.DataFrame([
    {
        "Model": r["model_name"],
        "R² Score": r["r2"],
        "RMSE": r["rmse"],
        "MAE": r["mae"],
        "MSE": r["mse"],
        "Explained Variance": r["explained_variance"],
        "Training Time (s)": r["training_time_sec"],
    }
    for r in results_list
])

# Rank by R2 descending, then RMSE ascending
comp_df = comp_df.sort_values(by=["R² Score", "RMSE"], ascending=[False, True]).reset_index(drop=True)

print("\nModel Comparison Table:")
print(comp_df.to_string(index=False))

# Select Best Model
best_res = sorted(results_list, key=lambda x: (x["r2"], -x["rmse"]), reverse=True)[0]
best_model_name = best_res["model_name"]
best_model_obj = best_res["model_obj"]

print(f"\n[BEST MODEL SELECTED]: {best_model_name}")
print(f"[BEST R² SCORE]    : {best_res['r2']:.4f}")
print(f"[LOWEST RMSE]      : {best_res['rmse']:.4f}")


# ==========================================================
# STEP 8 - Save Best Model & Metadata
# ==========================================================

print("\n" + "=" * 65)
print("STEP 8: Saving Best Model & Metadata...")
print("=" * 65)

best_model_path = MODEL_DIR / "best_model.pkl"
joblib.dump(best_model_obj, best_model_path)

meta_data = {
    "model_name": best_model_name,
    "task": "Productivity Prediction",
    "target_variable": "productivity_score",
    "metrics": {
        "r2_score": round(best_res["r2"], 4),
        "rmse": round(best_res["rmse"], 4),
        "mae": round(best_res["mae"], 4),
        "mse": round(best_res["mse"], 4),
        "explained_variance": round(best_res["explained_variance"], 4),
    },
    "feature_names": feature_names,
    "num_features": len(feature_names),
    "train_samples": len(X_train),
    "test_samples": len(X_test),
    "uses_raw_features": best_res["uses_raw_features"],
    "training_timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
}

meta_json_path = MODEL_DIR / "metadata.json"
model_meta_path = MODEL_DIR / "model_metadata.json"

with open(meta_json_path, "w", encoding="utf-8") as f:
    json.dump(meta_data, f, indent=4)

with open(model_meta_path, "w", encoding="utf-8") as f:
    json.dump(meta_data, f, indent=4)

print("Model & Metadata Saved Successfully!")
print(f"  - best_model.pkl     -> {best_model_path}")
print(f"  - metadata.json       -> {meta_json_path}")
print(f"  - model_metadata.json -> {model_meta_path}")


# ==========================================================
# STEP 9 - Generate Evaluation Artifacts & Plots
# ==========================================================

print("\n" + "=" * 65)
print("STEP 9: Generating Evaluation Plots & Metrics JSON...")
print("=" * 65)

print("Evaluation Started...")

# 1. Prediction Plot (Actual vs Predicted Scatter)
plt.figure(figsize=(8, 6))
plt.scatter(y_test, best_res["y_pred"], alpha=0.3, color="#2563eb", edgecolors="none", s=15)
plt.plot([0, 100], [0, 100], "--", color="#dc2626", linewidth=2, label="Ideal Fit (y = x)")
plt.title(f"Actual vs Predicted Productivity Score\n({best_model_name})", fontsize=12, fontweight="bold")
plt.xlabel("Actual Productivity Score", fontsize=10)
plt.ylabel("Predicted Productivity Score", fontsize=10)
plt.grid(True, linestyle="--", alpha=0.5)
plt.legend()
plt.tight_layout()
pred_plot_path = EVAL_DIR / "prediction_plot.png"
plt.savefig(pred_plot_path, dpi=300)
plt.close()
print(f"[+] Saved prediction_plot.png  -> {pred_plot_path}")

# 2. Residual Plot (Residual Distribution)
residuals = y_test - best_res["y_pred"]
plt.figure(figsize=(9, 4.5))
plt.subplot(1, 2, 1)
plt.scatter(best_res["y_pred"], residuals, alpha=0.3, color="#059669", s=15)
plt.axhline(0, color="#dc2626", linestyle="--", linewidth=1.5)
plt.title("Residuals vs Fitted Values", fontsize=11, fontweight="bold")
plt.xlabel("Predicted Productivity Score", fontsize=9)
plt.ylabel("Residuals (Actual - Pred)", fontsize=9)
plt.grid(True, linestyle="--", alpha=0.5)

plt.subplot(1, 2, 2)
plt.hist(residuals, bins=40, color="#10b981", edgecolor="black", alpha=0.7)
plt.axvline(0, color="#dc2626", linestyle="--", linewidth=1.5)
plt.title("Residuals Distribution", fontsize=11, fontweight="bold")
plt.xlabel("Residual Value", fontsize=9)
plt.ylabel("Frequency", fontsize=9)
plt.grid(True, linestyle="--", alpha=0.5)

plt.tight_layout()
residual_plot_path = EVAL_DIR / "residual_plot.png"
plt.savefig(residual_plot_path, dpi=300)
plt.close()
print(f"[+] Saved residual_plot.png    -> {residual_plot_path}")

# 3. Feature Importance Plot (Tree Model / Coefs)
plt.figure(figsize=(10, 6))
if hasattr(best_model_obj, "feature_importances_"):
    importances = best_model_obj.feature_importances_
    feat_imp = pd.Series(importances, index=feature_names).sort_values(ascending=True)
    feat_imp.plot(kind="barh", color="#3b82f6", edgecolor="black", alpha=0.85)
    plt.title(f"Feature Importances ({best_model_name})", fontsize=12, fontweight="bold")
    plt.xlabel("Importance Weight", fontsize=10)
elif hasattr(best_model_obj, "coef_"):
    coefs = pd.Series(best_model_obj.coef_, index=feature_names).sort_values(ascending=True)
    coefs.plot(kind="barh", color="#8b5cf6", edgecolor="black", alpha=0.85)
    plt.title(f"Feature Coefficients ({best_model_name})", fontsize=12, fontweight="bold")
    plt.xlabel("Coefficient Value", fontsize=10)
else:
    plt.text(0.5, 0.5, "Feature importances not available for this model", ha="center", va="center")

plt.grid(True, linestyle="--", alpha=0.5)
plt.tight_layout()
fi_plot_path = EVAL_DIR / "feature_importance.png"
plt.savefig(fi_plot_path, dpi=300)
plt.close()
print(f"[+] Saved feature_importance.png -> {fi_plot_path}")

# 4. Learning Curve
plt.figure(figsize=(8, 5))
# pyrefly: ignore [bad-unpacking]
train_sizes, train_scores, test_scores = learning_curve(
    best_model_obj,
    X_train_scaled if not best_res["uses_raw_features"] else X_train,
    y_train,
    cv=3,
    scoring="r2",
    train_sizes=np.linspace(0.1, 1.0, 5),
    n_jobs=-1,
)

train_mean = np.mean(train_scores, axis=1)
train_std = np.std(train_scores, axis=1)
test_mean = np.mean(test_scores, axis=1)
test_std = np.std(test_scores, axis=1)

plt.plot(train_sizes, train_mean, "o-", color="#2563eb", label="Training R² Score")
plt.plot(train_sizes, test_mean, "o-", color="#059669", label="Validation R² Score")
plt.fill_between(train_sizes, train_mean - train_std, train_mean + train_std, alpha=0.15, color="#2563eb")
plt.fill_between(train_sizes, test_mean - test_std, test_mean + test_std, alpha=0.15, color="#059669")

plt.title(f"Learning Curve ({best_model_name})", fontsize=12, fontweight="bold")
plt.xlabel("Training Set Size", fontsize=10)
plt.ylabel("R² Score", fontsize=10)
plt.grid(True, linestyle="--", alpha=0.5)
plt.legend(loc="lower right")
plt.tight_layout()
lc_plot_path = EVAL_DIR / "learning_curve.png"
plt.savefig(lc_plot_path, dpi=300)
plt.close()
print(f"[+] Saved learning_curve.png   -> {lc_plot_path}")

# 5. Save metrics.json
all_metrics_data = {
    "model_comparison": [
        {
            "model_name": r["model_name"],
            "r2_score": r["r2"],
            "rmse": r["rmse"],
            "mae": r["mae"],
            "mse": r["mse"],
            "explained_variance": r["explained_variance"],
            "training_time_sec": r["training_time_sec"],
        }
        for r in results_list
    ],
    "best_model": best_model_name,
}

eval_metrics_json_path = EVAL_DIR / "metrics.json"
with open(eval_metrics_json_path, "w", encoding="utf-8") as f:
    json.dump(all_metrics_data, f, indent=4)
print(f"[+] Saved metrics.json          -> {eval_metrics_json_path}")

print("Evaluation Completed.")


# ==========================================================
# STEP 10 - Generate PDF & DOCX Result Reports
# ==========================================================

print("\n" + "=" * 65)
print("STEP 10: Generating Reports (PDF & DOCX)...")
print("=" * 65)

# --- A. PDF REPORT GENERATOR (ReportLab) ---
def generate_productivity_pdf(pdf_path: Path):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            HRFlowable,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfgen import canvas

        class NumberedCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []

            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                num_pages = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self.drawString(54, 750, "EduPulse AI — Productivity Prediction ML Report")
                    self.setStrokeColor(colors.HexColor("#cbd5e1"))
                    self.setLineWidth(0.75)
                    self.line(54, 742, 612 - 54, 742)
                    page_str = f"Page {self._pageNumber} of {num_pages}"
                    self.drawRightString(612 - 54, 34, page_str)
                    self.drawString(54, 34, "EduPulse AI Final Year Project — Automated Execution Artifact")
                    super().showPage()
                super().save()

        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=72,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Heading1"],
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#1e3a8a"),
            fontName="Helvetica-Bold",
        )
        heading_style = ParagraphStyle(
            "DocHeading",
            parent=styles["Heading2"],
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#1e40af"),
            fontName="Helvetica-Bold",
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor("#334155"),
            fontName="Helvetica",
        )

        elements = []
        elements.append(Paragraph("Productivity Model Training & Evaluation Report", title_style))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y - %H:%M:%S')}", body_style))
        elements.append(Spacer(1, 10))
        elements.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2563eb"), spaceAfter=15))

        elements.append(Paragraph("1. Dataset Summary", heading_style))
        ds_info = [
            ["Metric", "Value"],
            ["Total Records", f"{len(df):,} records"],
            ["Features Count", f"{len(feature_names)} features"],
            ["Train / Test Split", f"80% ({len(X_train):,}) / 20% ({len(X_test):,})"],
            ["Target Variable", "productivity_score (0.00 to 100.00)"],
            ["Target Mean Score", f"{y.mean():.2f}"],
            ["Target Range", f"{y.min():.2f} to {y.max():.2f}"],
        ]
        t1 = Table(ds_info, colWidths=[200, 304])
        t1.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
        ]))
        elements.append(t1)
        elements.append(Spacer(1, 15))

        elements.append(Paragraph("2. Model Comparison Table", heading_style))
        comp_table_data = [["Model Name", "R² Score", "RMSE", "MAE", "Explained Var"]]
        for r in results_list:
            comp_table_data.append([
                r["model_name"],
                f"{r['r2']:.4f}",
                f"{r['rmse']:.4f}",
                f"{r['mae']:.4f}",
                f"{r['explained_variance']:.4f}",
            ])
        t2 = Table(comp_table_data, colWidths=[160, 85, 85, 85, 89])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0f766e')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f0fdf4')]),
        ]))
        elements.append(t2)
        elements.append(Spacer(1, 15))

        elements.append(Paragraph("3. Best Model Highlights", heading_style))
        best_info = [
            ["Attribute", "Value"],
            ["Selected Best Model", best_model_name],
            ["Highest R² Score", f"{best_res['r2']:.4f}"],
            ["Lowest RMSE", f"{best_res['rmse']:.4f}"],
            ["MAE", f"{best_res['mae']:.4f}"],
            ["Saved Model File", "models/productivity/best_model.pkl"],
            ["Saved Scaler File", "models/productivity/scaler.pkl"],
        ]
        t3 = Table(best_info, colWidths=[200, 304])
        t3.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#15803d')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f0fdf4')]),
        ]))
        elements.append(t3)

        doc.build(elements, canvasmaker=NumberedCanvas)
        print(f"[+] Saved productivity_result.pdf  -> {pdf_path}")
    except Exception as e:
        print(f"[!] PDF generation notice: {e}")

# --- B. DOCX REPORT GENERATOR (python-docx) ---
def generate_productivity_docx(docx_path: Path):
    try:
        # pyrefly: ignore [missing-import]
        from docx import Document
        # pyrefly: ignore [missing-import]
        from docx.shared import Pt, RGBColor

        doc = Document()
        
        # Title
        p_title = doc.add_paragraph()
        run_title = p_title.add_run("Productivity Model Training & Evaluation Report")
        run_title.font.size = Pt(20)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(30, 58, 138)
        
        p_sub = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y - %H:%M:%S')}")
        p_sub.runs[0].font.size = Pt(10)
        p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

        # 1. Dataset Summary
        h1 = doc.add_heading("1. Dataset Summary", level=2)
        h1.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        
        t1 = doc.add_table(rows=1, cols=2)
        hdr_cells1 = t1.rows[0].cells
        hdr_cells1[0].text = "Metric"
        hdr_cells1[1].text = "Value"
        
        ds_info = [
            ("Total Records", f"{len(df):,} records"),
            ("Features Count", f"{len(feature_names)} features"),
            ("Train / Test Split", f"80% ({len(X_train):,}) / 20% ({len(X_test):,})"),
            ("Target Variable", "productivity_score (0.00 to 100.00)"),
            ("Target Mean Score", f"{y.mean():.2f}"),
        ]
        for item, val in ds_info:
            row_cells = t1.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = val

        # 2. Model Comparison Table
        h2 = doc.add_heading("2. Model Comparison Table", level=2)
        h2.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        
        t2 = doc.add_table(rows=1, cols=5)
        hdr_cells2 = t2.rows[0].cells
        hdr_cells2[0].text = "Model Name"
        hdr_cells2[1].text = "R² Score"
        hdr_cells2[2].text = "RMSE"
        hdr_cells2[3].text = "MAE"
        hdr_cells2[4].text = "Explained Var"
        
        for r in results_list:
            row_cells = t2.add_row().cells
            row_cells[0].text = r["model_name"]
            row_cells[1].text = f"{r['r2']:.4f}"
            row_cells[2].text = f"{r['rmse']:.4f}"
            row_cells[3].text = f"{r['mae']:.4f}"
            row_cells[4].text = f"{r['explained_variance']:.4f}"

        # 3. Best Model Highlights
        h3 = doc.add_heading("3. Best Model Highlights", level=2)
        h3.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        
        p_best = doc.add_paragraph()
        p_best.add_run(f"Selected Best Model: {best_model_name}\n").bold = True
        p_best.add_run(f"Highest R² Score: {best_res['r2']:.4f}\n")
        p_best.add_run(f"Lowest RMSE: {best_res['rmse']:.4f}\n")
        p_best.add_run(f"Saved Path: models/productivity/best_model.pkl\n")

        doc.save(str(docx_path))
        print(f"[+] Saved productivity_result.docx -> {docx_path}")
    except Exception as e:
        print(f"[!] DOCX generation notice: {e}")

pdf_primary_path = REPORTS_DIR / "productivity_result.pdf"
docx_primary_path = REPORTS_DIR / "productivity_result.docx"

generate_productivity_pdf(pdf_primary_path)
generate_productivity_docx(docx_primary_path)

total_pipeline_time = time.time() - pipeline_start_time

print("\n" + "=" * 65)
print("TRAINING AND EVALUATION PIPELINE COMPLETED SUCCESSFULLY!")
print("=" * 65)
print(f"Best Model Selected : {best_model_name}")
print(f"Best Model R² Score : {best_res['r2']:.4f}")
print(f"Best Model RMSE     : {best_res['rmse']:.4f}")
print(f"Total Execution Time: {total_pipeline_time:.2f} seconds")
print("=" * 65)
