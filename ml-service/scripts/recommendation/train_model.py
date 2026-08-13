"""
train_model.py

EduPulse AI - AI Recommendation Engine ML Training & Evaluation Pipeline (Model 3)

Train multiple classification models (Decision Tree, Random Forest,
XGBoost / Gradient Boosting Classifier) using the recommendation dataset.
Generates evaluation metrics (Accuracy, Precision, Recall, Specificity, F1,
ROC AUC, Average Precision, MCC), confusion matrices, ROC/PR curves,
feature importances, learning curves, saves the best model & metadata,
and creates PDF & DOCX report documents.
"""

import os
import sys
import json
import time
import joblib
import numpy as np
import pandas as pd
from pathlib import Path
from datetime import datetime

# Configure stdout/stderr encoding to UTF-8 where possible
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

import matplotlib
matplotlib.use("Agg")  # Headless non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.model_selection import train_test_split, learning_curve
from sklearn.preprocessing import StandardScaler, label_binarize
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier

# Try importing XGBoost, fallback to GradientBoostingClassifier if unavailable
try:
    # pyrefly: ignore [missing-import]
    from xgboost import XGBClassifier
    HAS_XGB = True
except ImportError:
    from sklearn.ensemble import GradientBoostingClassifier
    HAS_XGB = False

from sklearn.metrics import (
    accuracy_score,
    precision_score,
    recall_score,
    f1_score,
    matthews_corrcoef,
    roc_auc_score,
    average_precision_score,
    confusion_matrix,
    classification_report,
    roc_curve,
    precision_recall_curve,
    auc,
)

# ==========================================================
# STEP 0 - DIRECTORY SETUP & PATH RESOLUTION
# ==========================================================
SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

if SCRIPT_DIR.name == "recommendation":
    BASE_DIR = SCRIPT_DIR.parent.parent
elif SCRIPT_DIR.name == "scripts":
    BASE_DIR = SCRIPT_DIR.parent
else:
    BASE_DIR = Path.cwd()

DATA_DIR = BASE_DIR / "data" / "recommendation"
MODEL_DIR = BASE_DIR / "models" / "recommendation"
EVAL_DIR = BASE_DIR / "evaluation" / "recommendation"
REPORTS_DIR = BASE_DIR / "reports"

DATA_DIR.mkdir(parents=True, exist_ok=True)
MODEL_DIR.mkdir(parents=True, exist_ok=True)
EVAL_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

pipeline_start_time = time.time()

# Target Class Mappings
CLASS_NAMES = [
    "Continue Current Skill",
    "Start Focus Session",
    "Take Short Break",
    "Practice Coding",
    "Revision",
    "Watch Learning Video",
    "Complete Pending Tasks",
    "Attempt Quiz",
]

CLASS_MAP = {i: name for i, name in enumerate(CLASS_NAMES)}


def main():
    # ==========================================================
    # STEP 1 - LOAD DATASET
    # ==========================================================
    print("=" * 65)
    print("STEP 1 - LOAD DATASET")
    print("=" * 65)

    dataset_path = DATA_DIR / "recommendation_dataset.csv"
    if not dataset_path.exists():
        raise FileNotFoundError(
            f"Dataset not found at '{dataset_path}'. Please run generate_dataset.py first."
        )

    df = pd.read_csv(dataset_path)

    print("Dataset Loaded Successfully")
    print(f"Dataset Path: {dataset_path}")
    print(f"Dataset Shape:\n{df.shape}")
    print("\nFirst 5 Records:")
    print(df.head())

    # ==========================================================
    # STEP 2 - PREPARE FEATURES
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 2 - PREPARE FEATURES")
    print("=" * 65)

    if "recommendation" not in df.columns:
        raise KeyError("Target column 'recommendation' not found in dataset!")

    feature_names = [col for col in df.columns if col != "recommendation"]
    X = df[feature_names]
    y = df["recommendation"].astype(int)

    unique_classes = np.sort(y.unique())
    print(f"Feature Names ({len(feature_names)}):\n{feature_names}")
    print(f"\nTarget Variable: recommendation")
    print(f"Target Classes ({len(unique_classes)}): {unique_classes.tolist()}")

    if len(unique_classes) != 8 or not np.array_equal(unique_classes, np.arange(8)):
        raise ValueError(
            f"Target must contain exactly 8 classes (0 to 7). Found: {unique_classes}"
        )

    print("\nClass Distribution:")
    class_dist = y.value_counts().sort_index()
    for cls_idx, count in class_dist.items():
        pct = (count / len(y)) * 100
        print(f"  Class {cls_idx} ({CLASS_NAMES[cls_idx]:<23}): {count:>6} ({pct:5.2f}%)")

    # ==========================================================
    # STEP 3 - TRAIN / TEST SPLIT
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 3 - TRAIN / TEST SPLIT")
    print("=" * 65)

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.20, random_state=42, stratify=y
    )

    # Save train and test split files
    train_df = pd.concat([X_train, y_train], axis=1)
    test_df = pd.concat([X_test, y_test], axis=1)

    train_file = DATA_DIR / "train.csv"
    test_file = DATA_DIR / "test.csv"

    train_df.to_csv(train_file, index=False)
    test_df.to_csv(test_file, index=False)

    print(f"Saved: {train_file}")
    print(f"Saved: {test_file}")
    print(f"\nTraining Samples: {len(X_train):,}")
    print(f"Testing Samples : {len(X_test):,}")
    print(f"X_train Shape   : {X_train.shape}")
    print(f"X_test Shape    : {X_test.shape}")
    print(f"y_train Shape   : {y_train.shape}")
    print(f"y_test Shape    : {y_test.shape}")

    # ==========================================================
    # STEP 4 - FEATURE SCALING
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 4 - FEATURE SCALING")
    print("=" * 65)

    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)

    scaler_file = MODEL_DIR / "scaler.pkl"
    joblib.dump(scaler, scaler_file)
    print(f"StandardScaler fitted and saved to: {scaler_file}")

    # ==========================================================
    # STEP 5 & STEP 6 - MODEL TRAINING & EVALUATION
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 5 - MODEL TRAINING")
    print("=" * 65)

    models = {}
    models["Decision Tree"] = DecisionTreeClassifier(random_state=42, max_depth=12)
    models["Random Forest"] = RandomForestClassifier(
        n_estimators=100, random_state=42, n_jobs=-1
    )

    if HAS_XGB:
        print("XGBoost detected: Using XGBClassifier")
        models["XGBoost"] = XGBClassifier(
            n_estimators=100,
            learning_rate=0.1,
            max_depth=6,
            random_state=42,
            eval_metric="mlogloss",
            objective="multi:softprob",
            num_class=8,
            n_jobs=-1,
        )
    else:
        print("XGBoost unavailable: Fallback to GradientBoostingClassifier")
        models["Gradient Boosting"] = GradientBoostingClassifier(
            n_estimators=100,
            learning_rate=0.1,
            max_depth=5,
            random_state=42,
        )

    results_list = []
    models_dict = {}

    y_test_bin = label_binarize(y_test, classes=range(8))

    for model_name, model in models.items():
        print(f"\nTraining {model_name}...")
        t_start = time.time()
        model.fit(X_train_scaled, y_train)
        t_train = time.time() - t_start

        print(f"Evaluating {model_name}...")
        y_pred = model.predict(X_test_scaled)
        y_proba = model.predict_proba(X_test_scaled)

        # Basic multi-class metrics
        acc = accuracy_score(y_test, y_pred)
        prec = precision_score(y_test, y_pred, average="weighted", zero_division=0)
        rec = recall_score(y_test, y_pred, average="weighted", zero_division=0)
        f1 = f1_score(y_test, y_pred, average="weighted", zero_division=0)
        mcc = matthews_corrcoef(y_test, y_pred)

        # Multi-class One-vs-Rest ROC AUC
        try:
            roc = roc_auc_score(y_test, y_proba, multi_class="ovr", average="weighted")
        except Exception:
            roc = 0.0

        # Multi-class One-vs-Rest Average Precision
        try:
            ap = average_precision_score(y_test_bin, y_proba, average="weighted")
        except Exception:
            ap = 0.0

        # Confusion Matrix
        cm = confusion_matrix(y_test, y_pred)

        # Calculate Multi-class Specificity using One-Vs-Rest
        class_specs = []
        for c in range(8):
            tp_c = cm[c, c]
            fp_c = cm[:, c].sum() - tp_c
            fn_c = cm[c, :].sum() - tp_c
            tn_c = cm.sum() - (tp_c + fp_c + fn_c)
            spec_c = (tn_c / (tn_c + fp_c)) if (tn_c + fp_c) > 0 else 0.0
            class_specs.append(spec_c)
        macro_spec = float(np.mean(class_specs))

        cls_report_str = classification_report(
            y_test, y_pred, target_names=CLASS_NAMES, zero_division=0
        )

        res = {
            "model_name": model_name,
            "model_obj": model,
            "accuracy": float(acc),
            "precision": float(prec),
            "recall": float(rec),
            "specificity": float(macro_spec),
            "f1": float(f1),
            "roc_auc": float(roc),
            "average_precision": float(ap),
            "mcc": float(mcc),
            "confusion_matrix": cm.tolist(),
            "training_time": float(t_train),
            "y_pred": y_pred,
            "y_proba": y_proba,
            "classification_report": cls_report_str,
        }

        results_list.append(res)
        models_dict[model_name] = res

    # ==========================================================
    # STEP 6 - PRINT MODEL EVALUATION SUMMARY
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 6 - MODEL EVALUATION")
    print("=" * 65)

    for r in results_list:
        print(f"\n--- {r['model_name']} ---")
        print(f"Accuracy         : {r['accuracy']:.4f}")
        print(f"Precision        : {r['precision']:.4f}")
        print(f"Recall           : {r['recall']:.4f}")
        print(f"Specificity      : {r['specificity']:.4f}")
        print(f"F1 Score         : {r['f1']:.4f}")
        print(f"ROC AUC          : {r['roc_auc']:.4f}")
        print(f"Average Precision: {r['average_precision']:.4f}")
        print(f"MCC              : {r['mcc']:.4f}")
        print(f"Training Time    : {r['training_time']:.2f}s")
        print("\nClassification Report:")
        print(r["classification_report"])

    # ==========================================================
    # STEP 7 - MODEL COMPARISON & SELECTION
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 7 - MODEL COMPARISON")
    print("=" * 65)

    comp_df = pd.DataFrame(
        [
            {
                "Model": r["model_name"],
                "Accuracy": r["accuracy"],
                "Precision": r["precision"],
                "Recall": r["recall"],
                "Specificity": r["specificity"],
                "F1 Score": r["f1"],
                "ROC AUC": r["roc_auc"],
                "Average Precision": r["average_precision"],
                "MCC": r["mcc"],
            }
            for r in results_list
        ]
    )

    print(comp_df.to_string(index=False))

    # Best model selection priority: 1. F1 Score, 2. MCC, 3. Accuracy
    sorted_results = sorted(
        results_list,
        key=lambda x: (x["f1"], x["mcc"], x["accuracy"]),
        reverse=True,
    )
    best_res = sorted_results[0]
    best_model_name = best_res["model_name"]
    best_model_obj = best_res["model_obj"]

    print("\nBest Model:")
    print(f"Name      : {best_model_name}")
    print(f"Best F1   : {best_res['f1']:.4f}")
    print(f"Best Acc  : {best_res['accuracy']:.4f}")
    print(f"Best MCC  : {best_res['mcc']:.4f}")

    # ==========================================================
    # STEP 8 - SAVE BEST MODEL & METADATA
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 8 - SAVE BEST MODEL")
    print("=" * 65)

    best_model_path = MODEL_DIR / "best_model.pkl"
    joblib.dump(best_model_obj, best_model_path)
    print(f"Saved Best Model: {best_model_path}")

    # Save metadata.json
    metadata = {
        "model_name": best_model_name,
        "target": "recommendation",
        "dataset_size": int(len(df)),
        "train_samples": int(len(X_train)),
        "test_samples": int(len(X_test)),
        "feature_names": feature_names,
        "class_names": CLASS_NAMES,
        "best_model": best_model_name,
        "metrics": {
            "accuracy": round(best_res["accuracy"], 4),
            "precision": round(best_res["precision"], 4),
            "recall": round(best_res["recall"], 4),
            "specificity": round(best_res["specificity"], 4),
            "f1": round(best_res["f1"], 4),
            "roc_auc": round(best_res["roc_auc"], 4),
            "average_precision": round(best_res["average_precision"], 4),
            "mcc": round(best_res["mcc"], 4),
        },
        "training_date": datetime.now().isoformat(),
    }

    metadata_path = MODEL_DIR / "metadata.json"
    with open(metadata_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=4)
    print(f"Saved Metadata: {metadata_path}")

    # Save model_metadata.json
    model_metadata = {
        "model_type": best_model_obj.__class__.__name__,
        "model_parameters": {
            k: str(v) for k, v in best_model_obj.get_params().items()
        },
        "scaler_type": "StandardScaler",
        "feature_names": feature_names,
        "target_classes": CLASS_NAMES,
        "best_model": best_model_name,
        "training_timestamp": datetime.now().isoformat(),
    }

    model_metadata_path = MODEL_DIR / "model_metadata.json"
    with open(model_metadata_path, "w", encoding="utf-8") as f:
        json.dump(model_metadata, f, indent=4)
    print(f"Saved Model Metadata: {model_metadata_path}")

    # ==========================================================
    # STEP 9 - GENERATE EVALUATION ARTIFACTS
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 9 - GENERATE EVALUATION")
    print("=" * 65)

    # 1. Confusion Matrix Plot
    fig, ax = plt.subplots(figsize=(10, 8))
    sns.heatmap(
        best_res["confusion_matrix"],
        annot=True,
        fmt="d",
        cmap="Blues",
        xticklabels=CLASS_NAMES,
        yticklabels=CLASS_NAMES,
        ax=ax,
    )
    ax.set_title(f"Confusion Matrix - {best_model_name}", fontsize=14, pad=12)
    ax.set_xlabel("Predicted Recommendation", fontsize=12)
    ax.set_ylabel("Actual Recommendation", fontsize=12)
    plt.xticks(rotation=45, ha="right", fontsize=9)
    plt.yticks(rotation=0, fontsize=9)
    plt.tight_layout()
    cm_path = EVAL_DIR / "confusion_matrix.png"
    plt.savefig(cm_path, dpi=300)
    plt.close()
    print(f"Saved: {cm_path}")

    # 2. Multi-Class ROC Curves
    fig, ax = plt.subplots(figsize=(10, 8))
    y_best_proba = best_res["y_proba"]
    for i in range(8):
        fpr_i, tpr_i, _ = roc_curve(y_test_bin[:, i], y_best_proba[:, i])
        auc_i = auc(fpr_i, tpr_i)
        ax.plot(
            fpr_i,
            tpr_i,
            lw=1.5,
            label=f"Class {i} ({CLASS_NAMES[i]}): AUC = {auc_i:.3f}",
        )

    ax.plot([0, 1], [0, 1], "k--", lw=1.5, label="Random Chance")
    ax.set_xlim([0.0, 1.0])
    ax.set_ylim([0.0, 1.05])
    ax.set_xlabel("False Positive Rate", fontsize=11)
    ax.set_ylabel("True Positive Rate", fontsize=11)
    ax.set_title(
        f"Multi-Class ROC Curves (One-vs-Rest) - {best_model_name}", fontsize=13
    )
    ax.legend(loc="lower right", fontsize=8)
    plt.tight_layout()
    roc_path = EVAL_DIR / "roc_curve.png"
    plt.savefig(roc_path, dpi=300)
    plt.close()
    print(f"Saved: {roc_path}")

    # 3. Multi-Class Precision-Recall Curves
    fig, ax = plt.subplots(figsize=(10, 8))
    for i in range(8):
        prec_i, rec_i, _ = precision_recall_curve(
            y_test_bin[:, i], y_best_proba[:, i]
        )
        ap_i = average_precision_score(y_test_bin[:, i], y_best_proba[:, i])
        ax.plot(
            rec_i,
            prec_i,
            lw=1.5,
            label=f"Class {i} ({CLASS_NAMES[i]}): AP = {ap_i:.3f}",
        )

    ax.set_xlim([0.0, 1.0])
    ax.set_ylim([0.0, 1.05])
    ax.set_xlabel("Recall", fontsize=11)
    ax.set_ylabel("Precision", fontsize=11)
    ax.set_title(
        f"Multi-Class Precision-Recall Curves (One-vs-Rest) - {best_model_name}",
        fontsize=13,
    )
    ax.legend(loc="lower left", fontsize=8)
    plt.tight_layout()
    pr_path = EVAL_DIR / "precision_recall_curve.png"
    plt.savefig(pr_path, dpi=300)
    plt.close()
    print(f"Saved: {pr_path}")

    # 4. Feature Importance Plot
    if hasattr(best_model_obj, "feature_importances_"):
        importances = best_model_obj.feature_importances_
        indices = np.argsort(importances)

        fig, ax = plt.subplots(figsize=(10, 8))
        ax.barh(range(len(indices)), importances[indices], align="center", color="#2563eb")
        ax.set_yticks(range(len(indices)))
        ax.set_yticklabels([feature_names[i] for i in indices], fontsize=9)
        ax.set_xlabel("Feature Importance Score", fontsize=11)
        ax.set_title(f"Feature Importance - {best_model_name}", fontsize=13)
        plt.tight_layout()
        fi_path = EVAL_DIR / "feature_importance.png"
        plt.savefig(fi_path, dpi=300)
        plt.close()
        print(f"Saved: {fi_path}")

    # 5. Learning Curve Plot
    # Subsample up to 5,000 samples for responsive, fast CV computation
    lc_samples = min(5000, len(X_train_scaled))
    X_lc = X_train_scaled[:lc_samples]
    y_lc = y_train.iloc[:lc_samples]

    # pyrefly: ignore [bad-unpacking]
    train_sizes, train_scores, test_scores = learning_curve(
        best_model_obj,
        X_lc,
        y_lc,
        cv=3,
        train_sizes=np.linspace(0.2, 1.0, 4),
        n_jobs=-1,
        scoring="accuracy",
    )

    train_mean = np.mean(train_scores, axis=1)
    train_std = np.std(train_scores, axis=1)
    test_mean = np.mean(test_scores, axis=1)
    test_std = np.std(test_scores, axis=1)

    fig, ax = plt.subplots(figsize=(9, 6))
    ax.plot(train_sizes, train_mean, "o-", color="#1d4ed8", label="Training Score")
    ax.plot(train_sizes, test_mean, "o-", color="#16a34a", label="Validation Score")
    ax.fill_between(
        train_sizes,
        train_mean - train_std,
        train_mean + train_std,
        alpha=0.15,
        color="#1d4ed8",
    )
    ax.fill_between(
        train_sizes,
        test_mean - test_std,
        test_mean + test_std,
        alpha=0.15,
        color="#16a34a",
    )
    ax.set_xlabel("Training Set Size", fontsize=11)
    ax.set_ylabel("Accuracy Score", fontsize=11)
    ax.set_title(f"Learning Curve - {best_model_name}", fontsize=13)
    ax.legend(loc="best", fontsize=10)
    ax.grid(True, linestyle="--", alpha=0.5)
    plt.tight_layout()
    lc_path = EVAL_DIR / "learning_curve.png"
    plt.savefig(lc_path, dpi=300)
    plt.close()
    print(f"Saved: {lc_path}")

    # 6. Classification Report Text File
    report_txt_path = EVAL_DIR / "classification_report.txt"
    with open(report_txt_path, "w", encoding="utf-8") as f:
        f.write(f"EDUPULSE AI - RECOMMENDATION MODEL ({best_model_name})\n")
        f.write("=" * 60 + "\n\n")
        f.write(best_res["classification_report"])
    print(f"Saved: {report_txt_path}")

    # 7. Comprehensive Metrics JSON
    metrics_all = {
        "dataset_size": int(len(df)),
        "train_samples": int(len(X_train)),
        "test_samples": int(len(X_test)),
        "feature_count": int(len(feature_names)),
        "feature_names": feature_names,
        "class_names": CLASS_NAMES,
        "best_model": best_model_name,
        "models": {},
    }

    for r in results_list:
        metrics_all["models"][r["model_name"]] = {
            "accuracy": round(r["accuracy"], 4),
            "precision": round(r["precision"], 4),
            "recall": round(r["recall"], 4),
            "specificity": round(r["specificity"], 4),
            "f1": round(r["f1"], 4),
            "roc_auc": round(r["roc_auc"], 4),
            "average_precision": round(r["average_precision"], 4),
            "mcc": round(r["mcc"], 4),
            "confusion_matrix": r["confusion_matrix"],
            "training_time": round(r["training_time"], 4),
        }

    metrics_json_path = EVAL_DIR / "metrics.json"
    with open(metrics_json_path, "w", encoding="utf-8") as f:
        json.dump(metrics_all, f, indent=4)
    print(f"Saved: {metrics_json_path}")

    # ==========================================================
    # STEP 10 - GENERATE REPORT (PDF & DOCX)
    # ==========================================================
    print("\n" + "=" * 65)
    print("STEP 10 - GENERATE REPORT")
    print("=" * 65)

    pdf_path = REPORTS_DIR / "recommendation_result.pdf"
    docx_path = REPORTS_DIR / "recommendation_result.docx"

    # --- A. PDF Report Generator (ReportLab) ---
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            HRFlowable,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfgen import canvas

        class NumberedCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []

            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                num_pages = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self.drawString(54, 750, "EduPulse AI — Recommendation Engine ML Report")
                    self.setStrokeColor(colors.HexColor("#cbd5e1"))
                    self.setLineWidth(0.75)
                    self.line(54, 742, 612 - 54, 742)
                    page_str = f"Page {self._pageNumber} of {num_pages}"
                    self.drawRightString(612 - 54, 34, page_str)
                    self.drawString(54, 34, "EduPulse AI Final Year Project — Automated Execution Artifact")
                    super().showPage()
                super().save()

        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            leftMargin=54,
            rightMargin=54,
            topMargin=72,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Heading1"],
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#1e3a8a"),
            fontName="Helvetica-Bold",
        )
        heading_style = ParagraphStyle(
            "DocHeading",
            parent=styles["Heading2"],
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1e40af"),
            fontName="Helvetica-Bold",
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#334155"),
            fontName="Helvetica",
        )

        elements = []
        elements.append(Paragraph("Recommendation Model Training & Evaluation Report", title_style))
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y - %H:%M:%S')}", body_style))
        elements.append(Spacer(1, 10))
        elements.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2563eb"), spaceAfter=12))

        # 1. Dataset Loaded Successfully & Summary
        elements.append(Paragraph("1. Dataset Summary", heading_style))
        ds_info = [
            ["Metric", "Value"],
            ["Dataset Loaded Successfully", "Yes (recommendation_dataset.csv)"],
            ["Dataset Shape", f"{df.shape[0]:,} rows × {df.shape[1]} columns"],
            ["Train / Test Split", f"80% ({len(X_train):,}) / 20% ({len(X_test):,})"],
            ["Target Column", "recommendation (8 Multi-class Target)"],
            ["Feature Scaling", "StandardScaler fitted on X_train"],
        ]
        t1 = Table(ds_info, colWidths=[200, 304])
        t1.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f8fafc')]),
        ]))
        elements.append(t1)
        elements.append(Spacer(1, 12))

        # 2. Model Comparison Table
        elements.append(Paragraph("2. Model Comparison Table", heading_style))
        comp_headers = ["Model", "Accuracy", "Precision", "Recall", "Specificity", "F1", "ROC AUC", "AP", "MCC"]
        comp_data = [comp_headers]
        for r in results_list:
            comp_data.append([
                r["model_name"],
                f"{r['accuracy']:.4f}",
                f"{r['precision']:.4f}",
                f"{r['recall']:.4f}",
                f"{r['specificity']:.4f}",
                f"{r['f1']:.4f}",
                f"{r['roc_auc']:.4f}",
                f"{r['average_precision']:.4f}",
                f"{r['mcc']:.4f}",
            ])
        t2 = Table(comp_data, colWidths=[100, 50, 50, 50, 55, 45, 50, 50, 44])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0f766e')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f0fdf4')]),
        ]))
        elements.append(t2)
        elements.append(Spacer(1, 12))

        # 3. Best Model Highlights & Saved Files
        elements.append(Paragraph("3. Best Model & Saved Files", heading_style))
        best_info = [
            ["Attribute", "Value"],
            ["Best Selected Model", best_model_name],
            ["Best F1 Score", f"{best_res['f1']:.4f}"],
            ["Best Accuracy", f"{best_res['accuracy']:.4f}"],
            ["Best MCC", f"{best_res['mcc']:.4f}"],
            ["Saved Model", "models/recommendation/best_model.pkl"],
            ["Saved Scaler", "models/recommendation/scaler.pkl"],
            ["Saved Metadata", "models/recommendation/metadata.json"],
            ["Saved Model Metadata", "models/recommendation/model_metadata.json"],
        ]
        t3 = Table(best_info, colWidths=[200, 304])
        t3.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#15803d')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f0fdf4')]),
        ]))
        elements.append(t3)

        doc.build(elements, canvasmaker=NumberedCanvas)
        print(f"Saved PDF Report -> {pdf_path}")
    except Exception as e:
        print(f"[!] PDF generation notice: {e}")

    # --- B. DOCX Report Generator (python-docx) ---
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        
        # Title
        p_title = doc.add_paragraph()
        run_title = p_title.add_run("Recommendation Model Training & Evaluation Report")
        run_title.font.size = Pt(18)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(30, 58, 138)
        
        p_sub = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y - %H:%M:%S')}")
        p_sub.runs[0].font.size = Pt(10)
        p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

        # 1. Dataset Summary
        h1 = doc.add_heading("1. Dataset Summary", level=2)
        h1.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        
        t1 = doc.add_table(rows=1, cols=2)
        hdr1 = t1.rows[0].cells
        hdr1[0].text = "Metric"
        hdr1[1].text = "Value"
        
        ds_info = [
            ("Dataset Loaded Successfully", "Yes (recommendation_dataset.csv)"),
            ("Dataset Shape", f"{df.shape[0]:,} rows × {df.shape[1]} columns"),
            ("Train / Test Split", f"80% ({len(X_train):,}) / 20% ({len(X_test):,})"),
            ("Target Column", "recommendation (8 Multi-class Target)"),
            ("Scaling", "StandardScaler"),
        ]
        for item, val in ds_info:
            row = t1.add_row().cells
            row[0].text = item
            row[1].text = val

        # 2. Model Comparison Table
        h2 = doc.add_heading("2. Model Comparison Table", level=2)
        h2.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        
        t2 = doc.add_table(rows=1, cols=9)
        hdr2 = t2.rows[0].cells
        headers = ["Model", "Accuracy", "Precision", "Recall", "Specificity", "F1", "ROC AUC", "AP", "MCC"]
        for idx, h_text in enumerate(headers):
            hdr2[idx].text = h_text
        
        for r in results_list:
            row = t2.add_row().cells
            row[0].text = r["model_name"]
            row[1].text = f"{r['accuracy']:.4f}"
            row[2].text = f"{r['precision']:.4f}"
            row[3].text = f"{r['recall']:.4f}"
            row[4].text = f"{r['specificity']:.4f}"
            row[5].text = f"{r['f1']:.4f}"
            row[6].text = f"{r['roc_auc']:.4f}"
            row[7].text = f"{r['average_precision']:.4f}"
            row[8].text = f"{r['mcc']:.4f}"

        # 3. Best Model & Saved Files
        h3 = doc.add_heading("3. Best Model & Saved Files", level=2)
        h3.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        
        p_best = doc.add_paragraph()
        p_best.add_run(f"Best Selected Model: {best_model_name}\n").bold = True
        p_best.add_run(f"Best F1 Score : {best_res['f1']:.4f}\n")
        p_best.add_run(f"Best Accuracy : {best_res['accuracy']:.4f}\n")
        p_best.add_run(f"Best MCC      : {best_res['mcc']:.4f}\n\n")
        p_best.add_run("Saved Files:\n").bold = True
        p_best.add_run(" - models/recommendation/best_model.pkl\n")
        p_best.add_run(" - models/recommendation/scaler.pkl\n")
        p_best.add_run(" - models/recommendation/metadata.json\n")
        p_best.add_run(" - models/recommendation/model_metadata.json\n")
        p_best.add_run(" - evaluation/recommendation/\n")
        p_best.add_run(" - reports/recommendation_result.pdf\n")
        p_best.add_run(" - reports/recommendation_result.docx\n")

        doc.save(str(docx_path))
        print(f"Saved DOCX Report -> {docx_path}")
    except Exception as e:
        print(f"[!] DOCX generation notice: {e}")

    # ==========================================================
    # FINAL COMPLETION OUTPUT
    # ==========================================================
    total_time = time.time() - pipeline_start_time
    print("\n" + "=" * 65)
    print("MODEL 3 TRAINING COMPLETED")
    print("=" * 65)
    print(f"\nBest Model:\n{best_model_name}")
    print(f"Best F1      : {best_res['f1']:.4f}")
    print(f"Best Accuracy: {best_res['accuracy']:.4f}")
    print(f"Best MCC     : {best_res['mcc']:.4f}")
    print(f"Total Execution Time: {total_time:.2f} seconds")
    print("\nAll artifacts saved successfully.")


if __name__ == "__main__":
    main()
