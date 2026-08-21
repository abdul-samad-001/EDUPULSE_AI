# -*- coding: utf-8 -*-
"""
EduPulse AI - Complete Technical Knowledge-Transfer Documentation Generator
Builds:
1. Complete_Project_Documentation.md
2. Complete_Project_Documentation.docx
3. Complete_Project_Documentation.pdf
"""

import os
import sys
import json
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = r"d:\FINAL YEAR\EDUPULSE_AI_NEW"
DOC_DIR = os.path.join(BASE_DIR, "documentation")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "Complete_Project_Documentation.md")
DOCX_PATH = os.path.join(DOC_DIR, "Complete_Project_Documentation.docx")
PDF_PATH = os.path.join(DOC_DIR, "Complete_Project_Documentation.pdf")

print(f"Target Directory: {DOC_DIR}")

# Load all text sections
from build_documentation_text import get_all_documentation_sections

sections = get_all_documentation_sections()

# 1. Write Markdown File
print("Writing Complete_Project_Documentation.md...")
full_md = "\n\n".join(sections)
with open(MD_PATH, "w", encoding="utf-8") as f:
    f.write(full_md)
print(f"[SUCCESS] Markdown created: {MD_PATH} ({len(full_md):,} characters, {os.path.getsize(MD_PATH):,} bytes)")

# 2. Build Word DOCX File
print("Building Complete_Project_Documentation.docx...")
doc = Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Configure Base Styles
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Dark slate

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    h.paragraph_format.space_before = Pt(14 if level == 1 else (10 if level == 2 else 6))
    h.paragraph_format.space_after = Pt(4)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Deep navy
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88) # Teal / Cyan accent
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate
    return h

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    # Give light gray background
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="0D9488"/></w:pBdr>')
    pPr.append(pBdr)
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>'))

def add_callout(doc, text, alert_type="NOTE"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    border_color = "0D9488" if alert_type == "NOTE" else ("F59E0B" if alert_type == "WARNING" else "EF4444")
    bg_color = "F0FDFA" if alert_type == "NOTE" else ("FFFBEB" if alert_type == "WARNING" else "FEF2F2")
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/></w:pBdr>')
    pPr.append(pBdr)
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>'))
    prefix = p.add_run(f"[{alert_type}] ")
    prefix.font.bold = True
    prefix.font.size = Pt(10)
    prefix.font.color.rgb = RGBColor(0x0D, 0x94, 0x88) if alert_type == "NOTE" else RGBColor(0xD9, 0x77, 0x06)
    body = p.add_run(text)
    body.font.size = Pt(10)

def parse_markdown_to_docx(doc, md_content):
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    in_table = False
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        # Parse table rows
        rows_data = []
        for line in t_lines:
            if re.match(r'^\s*\|?\s*[-:]+[-| :]*$', line):
                continue # separator
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if any(cells):
                rows_data.append(cells)
        if not rows_data:
            return
        
        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for r_idx, row_cells in enumerate(rows_data):
            for c_idx, text in enumerate(row_cells):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = text
                    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                    cell_p = cell.paragraphs[0]
                    cell_p.paragraph_format.space_before = Pt(2)
                    cell_p.paragraph_format.space_after = Pt(2)
                    if r_idx == 0:
                        set_cell_background(cell, "0F172A") # Navy header
                        for run in cell_p.runs:
                            run.font.bold = True
                            run.font.name = 'Arial'
                            run.font.size = Pt(9.5)
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        set_cell_background(cell, "F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
                        for run in cell_p.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(9.5)
                            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        # Space after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(4)

    while i < len(lines):
        line = lines[i]

        # Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle Tables
        if '|' in line and not line.startswith('>') and not line.startswith('#'):
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            flush_table(table_lines)
            table_lines = []

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            add_styled_heading(doc, line[2:].strip(), level=1)
        elif line.startswith('## '):
            add_styled_heading(doc, line[3:].strip(), level=2)
        elif line.startswith('### '):
            add_styled_heading(doc, line[4:].strip(), level=3)
        elif line.startswith('#### '):
            add_styled_heading(doc, line[5:].strip(), level=4)
        
        # Callouts / Blockquotes
        elif line.startswith('> [!IMPORTANT]') or line.startswith('> [!WARNING]') or line.startswith('> [!NOTE]'):
            alert_type = "IMPORTANT" if "IMPORTANT" in line else ("WARNING" if "WARNING" in line else "NOTE")
            # Collect following blockquote lines
            blockquote_text = []
            i += 1
            while i < len(lines) and lines[i].startswith('>'):
                cleaned = lines[i].lstrip('>').strip()
                if cleaned:
                    blockquote_text.append(cleaned)
                i += 1
            add_callout(doc, ' '.join(blockquote_text), alert_type)
            continue
        elif line.startswith('>'):
            add_callout(doc, line.lstrip('>').strip(), "NOTE")

        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            # Simple bold parsing
            parts = line[2:].split('**')
            for p_idx, part in enumerate(parts):
                run = p.add_run(part)
                if p_idx % 2 == 1:
                    run.font.bold = True
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            cleaned_text = re.sub(r'^\d+\.\s', '', line)
            parts = cleaned_text.split('**')
            for p_idx, part in enumerate(parts):
                run = p.add_run(part)
                if p_idx % 2 == 1:
                    run.font.bold = True

        # Regular Paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            parts = line.split('**')
            for p_idx, part in enumerate(parts):
                # Check for inline code
                subparts = part.split('`')
                for s_idx, subpart in enumerate(subparts):
                    run = p.add_run(subpart)
                    if p_idx % 2 == 1:
                        run.font.bold = True
                    if s_idx % 2 == 1:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
        
        i += 1

    if table_lines:
        flush_table(table_lines)

# Parse entire document into docx
print("Parsing sections into Word Document...")
parse_markdown_to_docx(doc, full_md)

doc.save(DOCX_PATH)
print(f"[SUCCESS] Word DOCX generated: {DOCX_PATH} ({os.path.getsize(DOCX_PATH):,} bytes)")

# 3. Build PDF File
print("Converting / Compiling Complete_Project_Documentation.pdf...")

# Method A: Win32 Word Automation (highest visual fidelity on Windows)
pdf_generated = False
try:
    import win32com.client
    import pythoncom
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    
    doc_abs = os.path.abspath(DOCX_PATH)
    pdf_abs = os.path.abspath(PDF_PATH)
    
    doc_com = word.Documents.Open(doc_abs)
    doc_com.SaveAs(pdf_abs, FileFormat=17) # wdFormatPDF = 17
    doc_com.Close()
    word.Quit()
    pdf_generated = True
    print(f"[SUCCESS] PDF generated via MS Word COM: {PDF_PATH} ({os.path.getsize(PDF_PATH):,} bytes)")
except Exception as e:
    print(f"[WARN] Win32 Word COM PDF export encountered: {e}. Falling back to ReportLab...")

# Method B: ReportLab Fallback
if not pdf_generated:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    pdf_doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=10
    )
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#0f172a"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#0d9488"),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=4
    )
    code_style = ParagraphStyle(
        'DocCode',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#0f172a"),
        backColor=colors.HexColor("#f1f5f9"),
        spaceBefore=4,
        spaceAfter=6
    )

    story = []
    
    for line in full_md.split('\n'):
        line_s = line.strip()
        if not line_s:
            story.append(Spacer(1, 4))
            continue
        if line_s.startswith('# '):
            story.append(Paragraph(line_s[2:].strip(), title_style))
        elif line_s.startswith('## '):
            story.append(Paragraph(line_s[3:].strip(), h1_style))
        elif line_s.startswith('### '):
            story.append(Paragraph(line_s[4:].strip(), h2_style))
        elif line_s.startswith('```'):
            continue
        elif '|' in line_s:
            story.append(Paragraph(line_s.replace('|', ' | '), body_style))
        else:
            clean_line = line_s.replace('**', '<b>').replace('`', '<code>')
            clean_line = re.sub(r'<b>(.*?)<b>', r'<b>\1</b>', clean_line)
            clean_line = re.sub(r'<code>(.*?)<code>', r'<code>\1</code>', clean_line)
            try:
                story.append(Paragraph(clean_line, body_style))
            except Exception:
                story.append(Paragraph(line_s, body_style))

    pdf_doc.build(story)
    print(f"[SUCCESS] PDF generated via ReportLab: {PDF_PATH} ({os.path.getsize(PDF_PATH):,} bytes)")

print("\n=================================================================")
print("  DOCUMENTATION COMPILATION COMPLETED SUCCESSFULLY")
print(f"  1. Markdown:  {MD_PATH}")
print(f"  2. Word DOCX: {DOCX_PATH}")
print(f"  3. PDF:       {PDF_PATH}")
print("=================================================================")
